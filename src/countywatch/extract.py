from __future__ import annotations

import csv
import io
import os
import re
import shutil
import warnings
import zipfile
from pathlib import Path
from urllib.parse import parse_qs, urlsplit

import fitz
import pytesseract
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning

from docx import Document
from PIL import Image

from .config import Settings
from .models import ExtractionResult
from .utils import normalize_space, normalize_text

warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
try:
    fitz.TOOLS.mupdf_display_errors(False)
    fitz.TOOLS.mupdf_display_warnings(False)
except Exception:
    pass


def _configure_tesseract() -> bool:
    found = shutil.which("tesseract")
    if found:
        pytesseract.pytesseract.tesseract_cmd = found
        return True
    candidates = [
        Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
        Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
    ]
    local_app_data = os.getenv("LOCALAPPDATA")
    if local_app_data:
        candidates.append(Path(local_app_data) / "Programs" / "Tesseract-OCR" / "tesseract.exe")
    for candidate in candidates:
        if candidate.exists():
            pytesseract.pytesseract.tesseract_cmd = str(candidate)
            return True
    return False


def extract_html(data: bytes) -> ExtractionResult:
    soup = BeautifulSoup(data, "lxml")
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "iframe"]):
        tag.decompose()
    for selector in ("nav", "header", "footer"):
        for tag in soup.select(selector):
            # Preserve unusually text-heavy elements because some vendor pages render agendas there.
            if len(normalize_space(tag.get_text(" "))) < 1500:
                tag.decompose()
    title = normalize_space(soup.title.get_text(" ") if soup.title else "")
    lines: list[str] = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "th", "caption", "time"]):
        value = normalize_space(element.get_text(" "))
        if value and (not lines or value != lines[-1]):
            lines.append(value)
    text = normalize_text("\n".join(lines))
    if len(text) < 300:
        text = normalize_text(soup.get_text("\n"))
    return ExtractionResult(text=text, method="html", metadata={"page_title": title})


def extract_pdf(data: bytes, settings: Settings) -> ExtractionResult:
    warnings: list[str] = []
    pages: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as document:
        page_count = document.page_count
        for page in document:
            pages.append(page.get_text("text", sort=True))
        text = normalize_text("\n\n".join(pages))
        meaningful = len(re.sub(r"\s+", "", text))
        needs_ocr = page_count > 0 and meaningful < max(350, page_count * 80)
        if needs_ocr and settings.enable_ocr:
            if _configure_tesseract():
                ocr_pages: list[str] = []
                limit = min(page_count, settings.max_ocr_pages)
                for index in range(limit):
                    page = document.load_page(index)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                    if pix.width * pix.height > 50_000_000:
                        warnings.append(f"OCR skipped oversized page {index + 1}")
                        continue
                    image = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_pages.append(pytesseract.image_to_string(image, lang="eng"))
                text = normalize_text("\n\n".join(ocr_pages))
                if page_count > limit:
                    warnings.append(f"OCR limited to the first {limit} of {page_count} pages")
                method = "pdf+ocr"
            else:
                warnings.append("PDF appears scanned, but Tesseract is not installed")
                method = "pdf"
        else:
            method = "pdf"
        metadata = dict(document.metadata or {})
    return ExtractionResult(text=text, method=method, page_count=page_count, metadata=metadata, warnings=warnings)


def extract_docx(data: bytes) -> ExtractionResult:
    # Office files are ZIP containers; reject pathological expansion before XML parsing.
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            expanded = sum(member.file_size for member in members)
            if len(members) > 5000 or expanded > 150_000_000:
                raise ValueError(
                    "DOCX archive exceeds safe expansion limits "
                    f"({len(members)} entries, {expanded} bytes)"
                )
    except zipfile.BadZipFile as exc:
        raise ValueError("Invalid DOCX archive") from exc
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        value = normalize_space(paragraph.text)
        if value:
            parts.append(value)
    for table in document.tables:
        for row in table.rows:
            values = [normalize_space(cell.text) for cell in row.cells]
            parts.append(" | ".join(value for value in values if value))
    properties = document.core_properties
    metadata = {
        "title": properties.title,
        "subject": properties.subject,
        "author": properties.author,
        "created": str(properties.created or ""),
        "modified": str(properties.modified or ""),
    }
    return ExtractionResult(text=normalize_text("\n".join(parts)), method="docx", metadata=metadata)


def extract_rtf(data: bytes) -> ExtractionResult:
    raw = data.decode("latin-1", errors="replace")
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", raw)
    raw = raw.replace("{", " ").replace("}", " ")
    return ExtractionResult(text=normalize_text(raw), method="rtf")


def extract_delimited(data: bytes, delimiter: str) -> ExtractionResult:
    decoded = data.decode("utf-8-sig", errors="replace")
    rows = csv.reader(io.StringIO(decoded), delimiter=delimiter)
    text = "\n".join(" | ".join(normalize_space(cell) for cell in row) for row in rows)
    return ExtractionResult(text=normalize_text(text), method="csv")


def extract_image(data: bytes, settings: Settings) -> ExtractionResult:
    if not settings.enable_ocr or not _configure_tesseract():
        return ExtractionResult(text="", method="image", warnings=["Image OCR unavailable"])
    image = Image.open(io.BytesIO(data))
    if image.width * image.height > 50_000_000:
        raise ValueError("Image exceeds safe OCR pixel limit")
    text = pytesseract.image_to_string(image, lang="eng")
    return ExtractionResult(text=normalize_text(text), method="image+ocr", page_count=1)


def youtube_video_id(url: str) -> str | None:
    split = urlsplit(url)
    if split.hostname in {"youtu.be", "www.youtu.be"}:
        return split.path.strip("/").split("/")[0] or None
    if split.hostname and "youtube.com" in split.hostname:
        if split.path == "/watch":
            return parse_qs(split.query).get("v", [None])[0]
        match = re.search(r"/(?:embed|live|shorts)/([^/?]+)", split.path)
        return match.group(1) if match else None
    return None


def extract_youtube(url: str, settings: Settings) -> ExtractionResult:
    if not settings.enable_youtube:
        return ExtractionResult(text="", method="youtube", warnings=["YouTube transcripts disabled"])
    video_id = youtube_video_id(url)
    if not video_id:
        return ExtractionResult(text="", method="youtube", warnings=["Could not identify YouTube video ID"])
    try:
        from youtube_transcript_api import YouTubeTranscriptApi

        api = YouTubeTranscriptApi()
        fetched = api.fetch(video_id, languages=["en", "en-US"])
        snippets = getattr(fetched, "snippets", fetched)
        lines = []
        for snippet in snippets:
            text = getattr(snippet, "text", None)
            if text is None and isinstance(snippet, dict):
                text = snippet.get("text")
            if text:
                lines.append(str(text))
        return ExtractionResult(text=normalize_text("\n".join(lines)), method="youtube-transcript")
    except Exception as exc:
        return ExtractionResult(text="", method="youtube", warnings=[f"Transcript unavailable: {exc}"])


def detect_kind(data: bytes, mime_type: str, url: str) -> str:
    mime = (mime_type or "").split(";", 1)[0].lower()
    suffix = Path(urlsplit(url).path).suffix.lower()
    if data.startswith(b"%PDF") or mime == "application/pdf" or suffix == ".pdf":
        return "pdf"
    if data.startswith(b"PK") and (mime.endswith("wordprocessingml.document") or suffix == ".docx"):
        return "docx"
    if mime in {"text/html", "application/xhtml+xml"} or suffix in {".html", ".htm"}:
        return "html"
    if mime.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return "image"
    if mime == "text/csv" or suffix == ".csv":
        return "csv"
    if suffix == ".rtf" or mime == "application/rtf":
        return "rtf"
    if mime.startswith("text/") or suffix == ".txt":
        return "text"
    # Some servers incorrectly return application/octet-stream for PDFs/HTML.
    prefix = data[:1000].lstrip().lower()
    if prefix.startswith((b"<!doctype html", b"<html", b"<?xml")):
        return "html"
    return "binary"


def extract_document(data: bytes, mime_type: str, url: str, settings: Settings) -> ExtractionResult:
    kind = detect_kind(data, mime_type, url)
    if kind == "pdf":
        return extract_pdf(data, settings)
    if kind == "docx":
        return extract_docx(data)
    if kind == "html":
        return extract_html(data)
    if kind == "image":
        return extract_image(data, settings)
    if kind == "csv":
        return extract_delimited(data, ",")
    if kind == "rtf":
        return extract_rtf(data)
    if kind == "text":
        return ExtractionResult(text=normalize_text(data.decode("utf-8-sig", errors="replace")), method="text")
    return ExtractionResult(text="", method="unsupported", warnings=[f"Unsupported document type: {mime_type or kind}"])
